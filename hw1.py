from docx import Document
import xlwt

doc=Document("待转文档.docx")
wb=xlwt.Workbook(encoding="utf-8")
ws=wb.add_sheet("问题导出")
list=[]
for i in range(len(doc.paragraphs)-2):
    list.append(doc.paragraphs[i].text)
ws.write(0,0,label="题目")
ws.write(0,1,label="选项1")
ws.write(0,2,label="选项2")
ws.write(0,3,label="选项3")
ws.write(0,4,label="选项4")
ws.write(0,5,label="答案")
del list[0]
list1=list[::6]
list2=list[1::6]
list3=list[2::6]
list4=list[3::6]
list5=list[4::6]

print(list2)
for i in range(len(list1)):
    ws.write(i+1,0,label=list1[i].split("？")[0].replace(str(i+1),"").replace(" ",""))
    ws.write(i+1,1,label=list2[i])
    ws.write(i+1,2,label=list3[i])
    ws.write(i+1,3,label=list4[i])
    ws.write(i+1,4,label=list5[i])
    ws.write(i+1,5,label=list1[i].split("（")[1].replace("）",""))

wb.save("转换表格.xls")